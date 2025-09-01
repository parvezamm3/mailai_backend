import io
import pandas as pd
from PyPDF2 import PdfReader # pip install pypdf2
import fitz
from docx import Document
from PIL import Image
import magic
import base64
import json
import asyncio
import requests
from config import Config
from pprint import pprint
import threading
# async def extract_text_from_attachment(file_bytes, filename):
#     """
#     Extracts plain text from various attachment file types.
#     """
#     file_extension = filename.split('.')[-1].lower()
#     print("File Extensions", file_extension)
    
#     if file_extension == 'txt':
#         print("File Type: Text")
#         return file_bytes.decode('utf-8', errors='ignore')
#     elif file_extension == 'pdf':
#         print("File Type: PDF")
#         try:
#             reader = PdfReader(io.BytesIO(file_bytes))
#             text = ""
#             for page in reader.pages:
#                 text += page.extract_text() or ""
#             return text
#         except Exception as e:
#             print(f"Error extracting text from PDF {filename}: {e}")
#             return None
#     elif file_extension == 'docx':
#         print("File Type: DOCX")
#         try:
#             document = Document(io.BytesIO(file_bytes))
#             return "\n".join([paragraph.text for paragraph in document.paragraphs])
#         except Exception as e:
#             print(f"Error extracting text from DOCX {filename}: {e}")
#             return None
#     elif file_extension == 'xlsx':
#         print("File Type: XLSX")
#         print(f"  Attempting to extract text from XLSX file: {filename}")
#         try:
#             # Read the Excel file into a pandas DataFrame
#             excel_data = pd.read_excel(io.BytesIO(file_bytes))
#             # Convert the entire DataFrame to a string for text extraction
#             text_content = excel_data.to_string(index=False, header=True)
#             return text_content
#         except Exception as e:
#             print(f"Error extracting text from XLSX {filename}: {e}")
#             return None

#     # NEW: Handler for CSV files
#     elif file_extension == 'csv':
#         print("File Type: CSV")
#         print(f"  Attempting to extract text from CSV file: {filename}")
#         try:
#             # Read the CSV file into a pandas DataFrame
#             csv_data = pd.read_csv(io.StringIO(file_bytes.decode('utf-8')))
#             # Convert the entire DataFrame to a string
#             text_content = csv_data.to_string(index=False, header=True)
#             return text_content
#         except Exception as e:
#             print(f"Error extracting text from CSV {filename}: {e}")
#             return None
    
#     # NEW: Image handling using Gemini API
#     elif file_extension in ['jpg', 'jpeg', 'png']:
#         print("File Type: Image")
#         try:
#             # === START OF NEW DEBUGGING CODE ===
#             # print(f"Debug: File size is {len(file_bytes)} bytes.")
#             # print(f"Debug: First 50 bytes of data: {file_bytes[:50]}")
#             # with open("temp_attachment.debug", "wb") as f:
#             #     f.write(file_bytes)
#             # print("Debug: Saved raw data to temp_attachment.debug for inspection.")
#             # === END OF NEW DEBUGGING CODE ===

#             # Add this crucial step to validate the image before sending to the API
#             # Image.open will raise an error if the file is not a valid image
#             # img = Image.open(io.BytesIO(file_bytes))
#             # img.verify()
#             # print(f"  Attempting to extract text from valid image {filename} using Gemini API...")
#             return await _extract_text_from_image_with_gemini(file_bytes, file_extension)
#         except Exception as e:
#             print(f"Error: The file {filename} is not a valid image. Skipping text extraction.")
#             print(f"Details: {e}")
#             return None
        
#     else:
#         print(f"Unsupported file type for text extraction: {filename}")
#         return None
    

async def extract_text_from_attachment(file_bytes, filename):
    """
    Extracts plain text from various attachment file types using asyncio.to_thread for blocking calls.
    """
    file_extension = filename.split('.')[-1].lower()
    print("File Extensions", file_extension)

    # Use a helper function for the synchronous work
    def _run_in_thread():
        # --- This is where all the synchronous, blocking code goes ---
        if file_extension == 'txt':
            return file_bytes.decode('utf-8', errors='ignore')
        elif file_extension == 'docx':
            document = Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in document.paragraphs)
        elif file_extension == 'xlsx':
            excel_data = pd.read_excel(io.BytesIO(file_bytes))
            return excel_data.to_string(index=False, header=True)
        elif file_extension == 'csv':
            csv_data = pd.read_csv(io.StringIO(file_bytes.decode('utf-8')))
            return csv_data.to_string(index=False, header=True)
        else:
            return None

    # Call the helper function in a separate thread
    if file_extension in ['txt', 'docx', 'xlsx', 'csv']:
        print(f"File Type: {file_extension.upper()}. Running in a separate thread...")
        try:
            return await asyncio.to_thread(_run_in_thread)
        except Exception as e:
            print(f"Error extracting text from {filename}: {e}")
            return None
    elif file_extension in ['pdf']:
        print("File Type: PDF")
        text_parts = []
        try:
            reader = await asyncio.to_thread(PdfReader(io.BytesIO(file_bytes)))
            text = "".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                print("Extracted Text", text)
                return text
        except Exception:
            pass
        print("File Type: PDF(File contains Image)")
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            pg_text = page.get_text("text")
            if pg_text.strip():
                text_parts.append(pg_text)
            else:
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.tobytes(output="png")
                ocr_text = await _extract_text_from_image_with_gemini(img_bytes, "png")
                if ocr_text:
                    text_parts.append(ocr_text)
            # print(text_parts)
            return "\n".join(text_parts)

    elif file_extension in ['jpg', 'jpeg', 'png']:
        print("File Type: Image")
        # This part of the code is already async and can be awaited directly
        try:
            return await _extract_text_from_image_with_gemini(file_bytes, file_extension)
        except Exception as e:
            print(f"Error extracting text from image {filename}: {e}")
            return None
    else:
        print(f"Unsupported file type for text extraction: {filename}")
        return None

def _get_mime_type(file_extension):
    """
    Helper to get the correct MIME type for an image file extension.
    """
    if file_extension in ['jpeg', 'jpg']:
        return 'image/jpeg'
    elif file_extension == 'png':
        return 'image/png'
    else:
        return 'application/octet-stream' 


async def _extract_text_from_image_with_gemini(image_bytes, file_extension):
    """
    Uses the Gemini API to perform OCR on an image and extract text.
    """
    try:
        # print(type(image_bytes))
        
        inferred_mime_type = magic.from_buffer(image_bytes, mime=True)
        # print(inferred_mime_type)
        base64_encoded_bytes = base64.b64encode(image_bytes)
        base64_encoded_string = base64_encoded_bytes.decode('utf-8')
        
        prompt = "Extract all text from the image. Do not add any extra commentary or formatting. Provide the raw text content."
        
        # Prepare the payload for the Gemini API call
        # mime_type =  _get_mime_type(file_extension)
        # if file_extension == 'jpg':
        #     mime_type = 'image/jpeg'
        payload = {
            "contents": [
                {
                    "role": "user",
                    "parts": [
                        { "text": prompt },
                        {
                            "inlineData": {
                                "mimeType": inferred_mime_type, # Assuming common image type, adjust if needed
                                "data": base64_encoded_string
                            }
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.1
            }
        }
        model = "gemini-2.0-flash-lite"
        
        apiUrl = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={Config.GEMINI_API_KEY}"
        

        headers = {'Content-Type': 'application/json'}
        response = await asyncio.to_thread(requests.post, apiUrl, headers=headers, json=payload)
        response.raise_for_status()
        response_data = response.json()
        # print("Image data extraction")
        
        
        try:
            # Check if 'candidates' exists and is not empty
            if response_data and response_data.get('candidates'):
                # print('Candidate Data')
                # pprint(response_data['candidates'][0]['content'])
                return response_data['candidates'][0]['content']['parts'][0]['text']
                
        except Exception as e:
            # This will catch any other unexpected errors during parsing
            print(f"An unexpected error occurred while parsing the Gemini API response: {e}")
    
    except Exception as e:
        print(f"Error using Gemini API for image text extraction: {e}")
        return None